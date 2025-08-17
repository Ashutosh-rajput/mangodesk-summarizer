import io
import mimetypes
import os
import shutil
import smtplib
import ssl
from email.message import EmailMessage

from PyPDF2 import PdfReader
from fastapi import UploadFile

from docx import Document as DocxDocument



def save_raw(file: UploadFile) -> str:
    base_dir = os.path.dirname(os.path.abspath(__file__))
    directory = os.path.join(base_dir, 'doc', 'raw')
    os.makedirs(directory, exist_ok=True)
    file_path = os.path.join(directory, file.filename)

    # Read the content from the file
    content = file.file.read()

    # Save the content to disk
    with open(file_path, 'wb') as out_file:
        out_file.write(content)

    # Extract text content based on file type
    content_type = file.content_type
    if content_type == "text/plain":
        return content.decode("utf-8")
    elif content_type == "application/pdf":
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = DocxDocument(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        raise Exception()

def get_raw() -> str:
    base_dir = os.path.dirname(os.path.abspath(__file__))

    directory = os.path.join(base_dir,'doc', 'raw')

    # --- Directory and File Validation ---
    if not os.path.exists(directory):
        return "Error: The 'srd_raw' directory was not found."
    if not os.path.isdir(directory):
        return f"Error: The path '{directory}' is not a directory."

    files = [f for f in os.listdir(directory) if os.path.isfile(os.path.join(directory, f))]
    if not files:
        return "Error: No files were found in the 'srd_raw' directory."

    # Sort files to ensure a consistent file is chosen every time
    files.sort()
    first_file_name = files[0]
    file_path = os.path.join(directory, first_file_name)

    # --- Content Extraction ---
    try:
        # Guess the file's MIME type to determine how to read it
        content_type, _ = mimetypes.guess_type(file_path)

        if content_type is None:
            raise Exception(f"Could not determine the file type for '{first_file_name}'.")

        # Read the file in binary mode ('rb') to handle any file type
        with open(file_path, 'rb') as file:
            content_bytes = file.read()

        # Extract text based on the detected content type
        if content_type == "text/plain":
            return content_bytes.decode("utf-8")

        elif content_type == "application/pdf":
            if not PdfReader:
                return "Error: PyPDF2 library is not installed. Please run 'pip install PyPDF2'."
            reader = PdfReader(io.BytesIO(content_bytes))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            if not DocxDocument:
                return "Error: python-docx library is not installed. Please run 'pip install python-docx'."
            doc = DocxDocument(io.BytesIO(content_bytes))
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)

        else:
            raise Exception(f"The file type '{content_type}' is not supported.")

    except Exception as e:
        return f"An error occurred while processing '{first_file_name}': {e}"



def delete_directory() -> None:
    base_dir = os.path.dirname(os.path.abspath(__file__))
    target_dir = os.path.join(base_dir, 'doc')

    if os.path.isdir(target_dir):
        shutil.rmtree(target_dir)


def save_summary(content: str):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    doc_dir = os.path.join(base_dir, 'doc')
    summary_file_path = os.path.join(doc_dir, 'summary.txt')

    try:
        # Create the target directory if it doesn't exist.
        # `exist_ok=True` prevents an error if the directory is already there.
        os.makedirs(doc_dir, exist_ok=True)

        # Open the file in 'write' mode ('w').
        # The 'with' statement ensures the file is automatically closed.
        with open(summary_file_path, 'w', encoding='utf-8') as f:
            f.write(content)

        print(f"Successfully saved summary to: {summary_file_path}")

    except IOError as e:
        # Handle potential I/O errors, like permission denied.
        print(f"Error: Could not save the file. {e}")
    except Exception as e:
        # Handle other unexpected errors.
        print(f"An unexpected error occurred: {e}")


# --- 3. The Method to Read the Summary ---

def read_summary() -> str:
    base_dir = os.path.dirname(os.path.abspath(__file__))
    doc_dir = os.path.join(base_dir, 'doc')
    summary_file_path = os.path.join(doc_dir, 'summary.txt')
    try:
        # Open the file in 'read' mode ('r').
        with open(summary_file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        print(f"Successfully read summary from: {summary_file_path}")
        return content

    except FileNotFoundError:
        # This is a common case, so we handle it specifically.
        print(f"Error: The file '{summary_file_path}' was not found.")
        return ""  # Return an empty string if file doesn't exist

    except IOError as e:
        print(f"Error: Could not read the file. {e}")
        return ""  # Return empty string on other I/O errors

    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return ""


async def send_summary_email(recipient_email: str):
    """
    Reads the summary from the file and emails it using Gmail.
    """
    # 1. Get credentials from environment variables
    sender_email = os.environ.get("GMAIL_SENDER_EMAIL")
    app_password = os.environ.get("GMAIL_APP_PASSWORD")

    if not sender_email or not app_password:
        # This is a server configuration error
        raise ValueError("GMAIL_SENDER_EMAIL and GMAIL_APP_PASSWORD environment variables must be set.")

    # 2. Read the summary content
    # This will raise FileNotFoundError if the file doesn't exist, which we'll catch in the endpoint
    summary_content = read_summary()

    # 3. Create the email message object
    msg = EmailMessage()
    msg['Subject'] = "Your Document Summary"
    msg['From'] = sender_email
    msg['To'] = recipient_email

    email_body = f"""
    Hello,

    Here is the document summary you requested:

    ---
    {summary_content}
    ---

    Generated by the Document Summarization API.
    """
    msg.set_content(email_body)

    # 4. Send the email using a secure connection
    context = ssl.create_default_context()
    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as smtp:
            smtp.login(sender_email, app_password)
            smtp.send_message(msg)
        print(f"Email sent successfully to {recipient_email}")
    except smtplib.SMTPAuthenticationError:
        # This error is specific to bad login credentials
        raise ConnectionRefusedError("Authentication failed. Check your GMAIL_SENDER_EMAIL and GMAIL_APP_PASSWORD.")
    except Exception as e:
        # Catch other potential errors (network, etc.)
        raise IOError(f"An error occurred while trying to send the email: {e}")
